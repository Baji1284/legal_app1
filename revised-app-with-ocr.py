# import streamlit as st
# import fitz  # PyMuPDF
# import docx
# import re
# import markdown
# from langchain.schema import HumanMessage, SystemMessage
# from langchain_core.prompts import ChatPromptTemplate
# import fitz  # PyMuPDF for extracting text from PDFs
# import docx  # For Word document support
# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain.schema import SystemMessage, HumanMessage
# import legal_prompt
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from google.generativeai import configure as google_configure
# import keys
# from io import BytesIO
# import markdown
# import re
# import report
# import io
# from PIL import Image
# import pytesseract #ocr
# from pdf2image import convert_from_bytes

# # # Configure Gemini
# #github
# # google_configure(api_key=st.secrets.GOOGLE_API_KEY)
# # gemini_model = ChatGoogleGenerativeAI(
# #     model="gemini-1.5-flash", 
# #     temperature=0  # deterministic behavior
# # )

# gemini_model = ChatGoogleGenerativeAI(
#     # model="gemini-1.5-flash",
#     model="gemini-1.5-pro" ,
#     # model="gemini-2.5-pro-preview-05-06",
#     google_api_key=keys.GOOGLE_API_KEY,
#     temperature=0.2,# deterministic behavior
#     max_output_tokens= 8192 
# )

# # Constants
# GEMINI_MAX_WORDS = 150000  # approximate large limit, adjust if needed

# #adding ocr for scanned docs

# # Extract text from PDF (including scanned with OCR)
# def extract_text_from_pdf(pdf_file):
#     text = ""
#     pdf_bytes = pdf_file.read()
    
#     # Try extracting selectable text using PyMuPDF
#     doc = fitz.open(stream=pdf_bytes, filetype="pdf")
#     for page in doc:
#         page_text = page.get_text("text")
#         text += page_text + "\n"
    
#     # If the text is mostly empty, apply OCR
#     if len(text.strip()) < 50:
#         st.info("Applying OCR for scanned contracts...")
#         # print("Applying OCR for scanned PDF...")
#         images = convert_from_bytes(pdf_bytes)
#         text = "\n".join([pytesseract.image_to_string(img) for img in images])
    
#     return text

# # Extract text from DOCX
# def extract_text_from_word(word_file):
#     doc = docx.Document(word_file)
#     text = "\n".join([para.text for para in doc.paragraphs])
    
#     # If no text found (possibly a scanned DOCX with images), apply OCR
#     if len(text.strip()) < 50:
#         st.info("Applying OCR for scanned contracts...")
#         # print("Applying OCR for scanned DOCX...")
#         for shape in doc.inline_shapes:
#             try:
#                 image_bytes = shape.image.blob
#                 image = Image.open(io.BytesIO(image_bytes))
#                 text += pytesseract.image_to_string(image) + "\n"
#             except Exception as e:
#                 print(f"OCR failed on shape: {e}")
    
#     return text

# # Chunk text based on word count
# def split_text(text, chunk_size):
#     words = text.split()
#     return [" ".join(words[i:i + chunk_size]) for i in range(0, len(words), chunk_size)]

# # Analyze contract using Gemini
# def analyze_contract(text, rulebook, instructions):
#     messages = [
#         SystemMessage(content=instructions),
#         SystemMessage(content=rulebook),
#         HumanMessage(content=f"Analyze the following contract text:\n{text}")
#     ]
#     response = gemini_model.invoke(messages)
#     return response.content if response else "No response from Gemini."

# # Default instructions (edit as needed)
# DEFAULT_INSTRUCTIONS = """
# You are a legal assistant specialized in Indian domestic contracts.
# You have to create a **Deviation Sheet** in the following tabular format as a response.

# ### Deviation Sheet Format (Mandatory):
# Column Names: 
# ●Original Clause Number and Reference 
# ●Original Clause
# ●Revised Clause
# ●Risk Summary

# | Original Clause Number and Reference | Original Clause             | Revised Clause                         | Risk Summary                                |
# |----------------------------------    |-----------------------------|-------------------------------------   |-------------------------------------------- |
# | Clause Number or Reference           | Original Clause & Reference | Modified clause(strikethrough and bold)| Explanation of risks and reasons for changes|


# **Rules for Deviation Sheet:**
# - DO NOT skip this table.
# - DO NOT include missing clauses (i.e., clauses required by the 'Guidelines for Clause Review' but absent from the contract) in the Deviation Sheet.
# - Only include in the Deviation Sheet those clauses that are present in the contract and have identified legal or commercial risks.
# - Only include the specific sub-clause (e.g., 17.2.3) where the legal or commercial risk exists.
# - DO NOT use placeholders such as “The entire Clause”or "Entire Clause" or “Refer Clause” in the deviation sheet,instead mention actual clause.
# - In “Revised Clause”, STRICTLY apply:
#     - ~~strikethrough~~ for deletions.
#     - **bold** for additions.
# - DO NOT paraphrase or improve clauses unless there is a clear legal risk.
# - Ensure that no changes are made to statements that convey the same meaning, even if they differ grammatically.
# - Ensure that all references, definitions, and clauses align with the contract’s existing terms and structure. 
# - Display all the modified clauses in the Deviation Sheet ONLY.

# """

# # Rulebook default (replace with your real rulebook content)
# DEFAULT_RULEBOOK = legal_prompt.rulebook

# # Streamlit UI
# st.set_page_config(page_title="LEGAL CONTRACT REVIEW: DOMESTIC ORDERS (INDIA)", layout="wide")

# # Sidebar
# with st.sidebar:
#     st.image("fm-logo.png", use_container_width=True)
#     st.write("Upload your Domestic Order Contracts (India) and receive an instant Legal Risk Assessment.")
#     uploaded_file = st.file_uploader("Upload your contract (PDF or Word)", type=["pdf", "docx"], key="contract")

# st.header("Legal Risk Assessment Tool — Version 1.0 - Revised (test)")

# # Editable Instructions
# st.subheader("Modify Default Instructions")
# instructions_text = st.text_area(
#     "Modify the default instructions as needed:", 
#     st.session_state.get("saved_instructions", DEFAULT_INSTRUCTIONS), 
#     height=300
# )
# if st.button("Save Instructions"):
#     st.session_state["saved_instructions"] = instructions_text
#     st.success("Instructions saved.")

# # Editable Rulebook
# st.subheader("Modify Rulebook")
# rulebook_text = st.text_area(
#     "Modify the rulebook as needed:", 
#     st.session_state.get("saved_rulebook", DEFAULT_RULEBOOK), 
#     height=500
# )
# if st.button("Save Rulebook"):
#     st.session_state["saved_rulebook"] = rulebook_text
#     st.success("Rulebook saved.")

# # Load saved instruction/rulebook
# instructions_text = st.session_state.get("saved_instructions", DEFAULT_INSTRUCTIONS)
# rulebook_text = st.session_state.get("saved_rulebook", DEFAULT_RULEBOOK)

# # File processing
# if uploaded_file and rulebook_text and instructions_text:
#     with st.spinner("Extracting text from contract..."):
#         file_type = uploaded_file.name.split('.')[-1].lower()
#         if file_type == "pdf":
#             contract_text = extract_text_from_pdf(uploaded_file)
#         elif file_type == "docx":
#             contract_text = extract_text_from_word(uploaded_file)
#         else:
#             st.error("Unsupported file type.")
#             st.stop()

#     text_chunks = split_text(contract_text, GEMINI_MAX_WORDS)

#     if st.button("Analyze Risks"):
#         with st.spinner("Analyzing contract risks..."):
#             risk_analysis = ""
#             for chunk in text_chunks:
#                 result = analyze_contract(chunk, rulebook_text, instructions_text)
#                 result = re.sub(r"~~(.*?)~~", r"<del>\1</del>", result)  # markdown strikethrough
#                 result = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", result)  # markdown bold
#                 risk_analysis += result + "\n\n"

#             st.subheader("Risk Analysis Report")
#             st.markdown(risk_analysis, unsafe_allow_html=True)

#             # Convert Markdown to HTML
#             risk_analysis_html = markdown.markdown(risk_analysis, extensions=["tables"])
#             docx_bytes = report.report_downloader(risk_analysis_html, logo_path="fm-logo.png")

#             st.download_button(
#                 label="Download Report as Word (.docx)",
#                 data=docx_bytes,
#                 file_name="Risk_Analysis_Report.docx",
#                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#             )
# else:
#     st.info("Please upload a contract and ensure the rulebook + instructions are filled in.")









##########################################new approach############################################

import streamlit as st
import fitz  # PyMuPDF
import docx
import re
import markdown
from langchain.schema import HumanMessage, SystemMessage
from langchain_core.prompts import ChatPromptTemplate
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from google.generativeai import configure as google_configure
import keys
from io import BytesIO
import io
from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes
import pandas as pd
import numpy as np
import cv2
import legal_prompt
import report

# Configure Gemini
gemini_model = ChatGoogleGenerativeAI(
    model="gemini-1.5-pro",
    google_api_key=keys.GOOGLE_API_KEY,
    temperature=0.2,
    max_output_tokens=8192
)

# Constants
GEMINI_MAX_WORDS = 150000  # approximate large limit

# OCR Configuration for better accuracy
def configure_ocr():
    """Configure OCR settings for better accuracy"""
    custom_config = r'--oem 3 --psm 6 -l eng'
    return custom_config

# Image preprocessing for better OCR
def preprocess_image_for_ocr(image):
    """Preprocess image to improve OCR accuracy"""
    # Convert PIL Image to numpy array
    img_array = np.array(image)
    
    # Convert to grayscale if not already
    if len(img_array.shape) == 3:
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
    else:
        gray = img_array
    
    # Apply thresholding to get better contrast
    _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    
    # Denoise
    denoised = cv2.medianBlur(thresh, 1)
    
    # Convert back to PIL Image
    return Image.fromarray(denoised)

# Extract text from tables in PDF
def extract_tables_from_pdf(pdf_bytes):
    """Extract tables from PDF using multiple methods"""
    tables_text = []
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    
    for page_num, page in enumerate(doc):
        # Method 1: Try to extract tables using PyMuPDF
        tabs = page.find_tables()
        if tabs:
            for tab in tabs:
                try:
                    # Extract table as pandas DataFrame
                    df = tab.to_pandas()
                    # Convert DataFrame to text format
                    table_text = f"\n[Table from Page {page_num + 1}]\n"
                    table_text += df.to_string(index=False)
                    tables_text.append(table_text)
                except:
                    # If pandas conversion fails, try extracting raw
                    extracted = tab.extract()
                    table_text = f"\n[Table from Page {page_num + 1}]\n"
                    for row in extracted:
                        row_text = " | ".join([str(cell) if cell else "" for cell in row])
                        table_text += row_text + "\n"
                    tables_text.append(table_text)
    
    doc.close()
    return "\n".join(tables_text)

# Enhanced PDF text extraction with OCR and table support
def extract_text_from_pdf(pdf_file):
    """Extract text from PDF including scanned documents and tables"""
    text = ""
    tables_text = ""
    pdf_bytes = pdf_file.read()
    
    # First, try to extract tables
    try:
        tables_text = extract_tables_from_pdf(pdf_bytes)
    except Exception as e:
        st.warning(f"Table extraction encountered an issue: {e}")
    
    # Extract regular text using PyMuPDF
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page_num, page in enumerate(doc):
        page_text = page.get_text("text")
        text += f"\n--- Page {page_num + 1} ---\n"
        text += page_text + "\n"
    doc.close()
    
    # Check if we need OCR (if extracted text is too short)
    if len(text.strip()) < 100:  # Threshold for detecting scanned documents
        st.info("📄 Detected scanned document. Applying OCR for text extraction...")
        
        # Convert PDF to images and apply OCR
        try:
            images = convert_from_bytes(pdf_bytes, dpi=300)  # Higher DPI for better OCR
            ocr_config = configure_ocr()
            
            for i, img in enumerate(images):
                st.progress((i + 1) / len(images), text=f"Processing page {i + 1} of {len(images)}...")
                
                # Preprocess image for better OCR
                processed_img = preprocess_image_for_ocr(img)
                
                # Apply OCR
                page_text = pytesseract.image_to_string(processed_img, config=ocr_config)
                text += f"\n--- Page {i + 1} (OCR) ---\n"
                text += page_text + "\n"
        except Exception as e:
            st.error(f"OCR processing failed: {e}")
    
    # Combine regular text with tables
    if tables_text:
        text += "\n\n--- EXTRACTED TABLES ---\n" + tables_text
    
    return text

# Extract tables from Word documents
def extract_tables_from_word(doc):
    """Extract tables from Word document"""
    tables_text = []
    
    for i, table in enumerate(doc.tables):
        table_text = f"\n[Table {i + 1}]\n"
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_text += " | ".join(row_data) + "\n"
        tables_text.append(table_text)
    
    return "\n".join(tables_text)

# Enhanced Word document text extraction
def extract_text_from_word(word_file):
    """Extract text from Word documents including tables and embedded images"""
    doc = docx.Document(word_file)
    text = ""
    
    # Extract regular paragraph text
    paragraphs_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    text += paragraphs_text
    
    # Extract tables
    tables_text = extract_tables_from_word(doc)
    if tables_text:
        text += "\n\n--- EXTRACTED TABLES ---\n" + tables_text
    
    # Check if document might be scanned (very little text extracted)
    if len(text.strip()) < 100:
        st.info("📄 Detected potentially scanned content in Word document. Applying OCR...")
        
        # Extract and OCR embedded images
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_bytes = rel.target_part.blob
                    image = Image.open(io.BytesIO(image_bytes))
                    
                    # Preprocess and apply OCR
                    processed_img = preprocess_image_for_ocr(image)
                    ocr_config = configure_ocr()
                    ocr_text = pytesseract.image_to_string(processed_img, config=ocr_config)
                    
                    text += f"\n--- Extracted from embedded image ---\n"
                    text += ocr_text + "\n"
                except Exception as e:
                    st.warning(f"Failed to process embedded image: {e}")
    
    return text

# Extract text from image files
def extract_text_from_image(image_file):
    """Extract text from image files (PNG, JPG, JPEG, TIFF, BMP)"""
    st.info("🖼️ Processing image file with OCR...")
    
    try:
        # Open image
        image = Image.open(image_file)
        
        # Preprocess image for better OCR
        processed_img = preprocess_image_for_ocr(image)
        
        # Apply OCR with custom configuration
        ocr_config = configure_ocr()
        text = pytesseract.image_to_string(processed_img, config=ocr_config)
        
        # Also try to detect tables in the image (basic approach)
        # This is a simplified version - for production, consider using specialized table detection
        text_with_layout = pytesseract.image_to_string(processed_img, config=ocr_config + ' --psm 6')
        
        if text_with_layout != text:
            text += "\n\n--- Layout preserved text ---\n" + text_with_layout
        
        return text
    except Exception as e:
        st.error(f"Failed to extract text from image: {e}")
        return ""

# Smart text extraction based on file type
def extract_text_from_file(uploaded_file):
    """Smart extraction based on file type with comprehensive support"""
    file_type = uploaded_file.name.split('.')[-1].lower()
    
    # Reset file pointer
    uploaded_file.seek(0)
    
    if file_type == "pdf":
        return extract_text_from_pdf(uploaded_file)
    elif file_type in ["docx", "doc"]:
        return extract_text_from_word(uploaded_file)
    elif file_type in ["png", "jpg", "jpeg", "tiff", "bmp"]:
        return extract_text_from_image(uploaded_file)
    else:
        st.error(f"Unsupported file type: {file_type}")
        return ""

# Enhanced text chunking that preserves table structure
def split_text(text, chunk_size):
    """Split text while trying to preserve table and section integrity"""
    # First, try to identify tables and keep them together
    lines = text.split('\n')
    chunks = []
    current_chunk = []
    current_word_count = 0
    in_table = False
    
    for line in lines:
        # Check if line is part of a table (simple heuristic)
        if '[Table' in line or '|' in line or '---' in line:
            in_table = True
        elif in_table and line.strip() == "":
            in_table = False
        
        words_in_line = len(line.split())
        
        # If adding this line would exceed chunk size and we're not in a table
        if current_word_count + words_in_line > chunk_size and not in_table:
            if current_chunk:
                chunks.append('\n'.join(current_chunk))
            current_chunk = [line]
            current_word_count = words_in_line
        else:
            current_chunk.append(line)
            current_word_count += words_in_line
            
            # If we're in a table and approaching limit, finish after table
            if in_table and current_word_count > chunk_size * 0.8:
                if not line.strip() or not ('|' in line or '---' in line):
                    chunks.append('\n'.join(current_chunk))
                    current_chunk = []
                    current_word_count = 0
                    in_table = False
    
    # Add remaining chunk
    if current_chunk:
        chunks.append('\n'.join(current_chunk))
    
    return chunks if chunks else [text]

# Analyze contract using Gemini
def analyze_contract(text, rulebook, instructions):
    """Analyze contract with enhanced context awareness"""
    messages = [
        SystemMessage(content=instructions),
        SystemMessage(content=rulebook),
        HumanMessage(content=f"Analyze the following contract text (including any tables):\n{text}")
    ]
    response = gemini_model.invoke(messages)
    return response.content if response else "No response from Gemini."

# Default instructions
DEFAULT_INSTRUCTIONS = """
You are a legal assistant specialized in Indian domestic contracts.
You have to create a **Deviation Sheet** in the following tabular format as a response.

### Deviation Sheet Format (Mandatory):
Column Names: 
●Original Clause Number and Reference 
●Original Clause
●Revised Clause
●Risk Summary

| Original Clause Number and Reference | Original Clause             | Revised Clause                         | Risk Summary                                |
|----------------------------------    |-----------------------------|-------------------------------------   |-------------------------------------------- |
| Clause Number or Reference           | Original Clause & Reference | Modified clause(strikethrough and bold)| Explanation of risks and reasons for changes|

**Rules for Deviation Sheet:**
- DO NOT skip this table.
- DO NOT include missing clauses (i.e., clauses required by the 'Guidelines for Clause Review' but absent from the contract) in the Deviation Sheet.
- Only include in the Deviation Sheet those clauses that are present in the contract and have identified legal or commercial risks.
- Only include the specific sub-clause (e.g., 17.2.3) where the legal or commercial risk exists.
- DO NOT use placeholders such as "The entire Clause"or "Entire Clause" or "Refer Clause" in the deviation sheet,instead mention actual clause.
- In "Revised Clause", STRICTLY apply:
    - ~~strikethrough~~ for deletions.
    - **bold** for additions.
- DO NOT paraphrase or improve clauses unless there is a clear legal risk.
- Ensure that no changes are made to statements that convey the same meaning, even if they differ grammatically.
- Ensure that all references, definitions, and clauses align with the contract's existing terms and structure. 
- Display all the modified clauses in the Deviation Sheet ONLY.
- Pay special attention to data in tables as they often contain critical terms and conditions.
"""

DEFAULT_RULEBOOK = legal_prompt.rulebook

# Streamlit UI
st.set_page_config(page_title="LEGAL CONTRACT REVIEW: DOMESTIC ORDERS (INDIA)", layout="wide")

# Sidebar
with st.sidebar:
    st.image("fm-logo.png", use_container_width=True)
    st.write("Upload your Domestic Order Contracts (India) and receive an instant Legal Risk Assessment.")
    st.write("**Supported formats:** PDF, Word (.docx), Images (PNG, JPG, JPEG, TIFF, BMP)")
    st.write("**Features:**")
    st.write("✅ OCR for scanned documents")
    st.write("✅ Table extraction")
    st.write("✅ Multi-page support")
    
    uploaded_file = st.file_uploader(
        "Upload your contract", 
        type=["pdf", "docx", "doc", "png", "jpg", "jpeg", "tiff", "bmp"], 
        key="contract"
    )

st.header("Legal Risk Assessment Tool — Version 2.0 - Enhanced OCR & Table Support")

# Editable Instructions
with st.expander("⚙️ Modify Default Instructions", expanded=False):
    instructions_text = st.text_area(
        "Modify the default instructions as needed:", 
        st.session_state.get("saved_instructions", DEFAULT_INSTRUCTIONS), 
        height=300
    )
    if st.button("Save Instructions"):
        st.session_state["saved_instructions"] = instructions_text
        st.success("Instructions saved.")

# Editable Rulebook
with st.expander("📚 Modify Rulebook", expanded=False):
    rulebook_text = st.text_area(
        "Modify the rulebook as needed:", 
        st.session_state.get("saved_rulebook", DEFAULT_RULEBOOK), 
        height=500
    )
    if st.button("Save Rulebook"):
        st.session_state["saved_rulebook"] = rulebook_text
        st.success("Rulebook saved.")

# Load saved instruction/rulebook
instructions_text = st.session_state.get("saved_instructions", DEFAULT_INSTRUCTIONS)
rulebook_text = st.session_state.get("saved_rulebook", DEFAULT_RULEBOOK)

# File processing
if uploaded_file and rulebook_text and instructions_text:
    # Display file info
    file_info = st.container()
    with file_info:
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("File Name", uploaded_file.name)
        with col2:
            st.metric("File Size", f"{uploaded_file.size / 1024:.2f} KB")
        with col3:
            st.metric("File Type", uploaded_file.name.split('.')[-1].upper())
    
    # Extract text button
    if st.button("🔍 Extract & Analyze Contract", type="primary"):
        with st.spinner("Extracting text from contract..."):
            contract_text = extract_text_from_file(uploaded_file)
            
            if not contract_text or len(contract_text.strip()) < 10:
                st.error("Failed to extract meaningful text from the document. Please ensure the file is not corrupted.")
                st.stop()
            
            # Show extraction summary
            st.success(f"✅ Successfully extracted {len(contract_text.split())} words from the document")
            
            # Show preview of extracted text
            with st.expander("Preview Extracted Text (first 1000 characters)", expanded=False):
                st.text(contract_text[:1000] + "..." if len(contract_text) > 1000 else contract_text)
            
            # Check if tables were found
            if "--- EXTRACTED TABLES ---" in contract_text:
                st.info("📊 Tables detected and extracted from the document")
        
        # Chunk the text
        text_chunks = split_text(contract_text, GEMINI_MAX_WORDS)
        st.info(f"Document split into {len(text_chunks)} chunk(s) for analysis")
        
        # Analyze the contract
        with st.spinner("🤖 Analyzing contract risks with AI..."):
            risk_analysis = ""
            progress_bar = st.progress(0)
            
            for i, chunk in enumerate(text_chunks):
                progress_bar.progress((i + 1) / len(text_chunks), 
                                     text=f"Analyzing chunk {i + 1} of {len(text_chunks)}...")
                
                result = analyze_contract(chunk, rulebook_text, instructions_text)
                
                # Format the result
                result = re.sub(r"~~(.*?)~~", r"<del>\1</del>", result)
                result = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", result)
                risk_analysis += result + "\n\n"
            
            progress_bar.empty()
            
            # Display results
            st.subheader("📋 Risk Analysis Report")
            st.markdown(risk_analysis, unsafe_allow_html=True)
            
            # Generate downloadable report
            risk_analysis_html = markdown.markdown(risk_analysis, extensions=["tables"])
            docx_bytes = report.report_downloader(risk_analysis_html, logo_path="fm-logo.png")
            
            st.download_button(
                label="📥 Download Report as Word (.docx)",
                data=docx_bytes,
                file_name=f"Risk_Analysis_Report_{uploaded_file.name.split('.')[0]}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
else:
    # Instructions when no file is uploaded
    st.info("👆 Please upload a contract document to begin analysis")
    
    with st.expander("ℹ️ How to use this tool"):
        st.write("""
        1. **Upload a contract document** using the sidebar (PDF, Word, or Image format)
        2. **Review/modify instructions and rulebook** if needed (optional)
        3. **Click "Extract & Analyze Contract"** to process the document
        4. **Review the risk analysis** displayed on screen
        5. **Download the report** as a Word document for further review
        
        **Features:**
        - Automatic OCR for scanned documents
        - Table extraction and analysis
        - Multi-format support
        - Intelligent text chunking for large documents
        """)



